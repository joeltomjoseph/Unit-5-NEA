''' General functions that are used throughout the program. '''
import os
import sys
import platform
import subprocess
import shutil
import smtplib, ssl
import random
from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE

if platform.system() == "Windows": # If the system is Windows, get the desktop directory from the USERPROFILE environment variable
    desktop = os.path.join(os.path.join(os.environ['USERPROFILE']), 'Desktop') # Get the desktop directory
else: # If the system is not Windows (MacOS or Linux), get the desktop directory from the home directory
    desktop = os.path.join(os.path.join(os.path.expanduser('~')), 'Desktop')

# Adapted from https://stackoverflow.com/questions/17576366/print-out-the-whole-directory-tree
def getDirectoryStructure(rootDir) -> dict:
    ''' Recursively get the directory structure of a given directory.
    Returns a nested dictionary of all folders and a list files within each folder, within a given directory.
    Ignores hidden files and folders (starting with a .) '''
    items = {}
    # First, get all folders and create their nested dictionaries
    for item in os.listdir(rootDir):
        filePath = os.path.join(rootDir, item) 
        #print(item)

        if item[0] != '.':
            if os.path.isdir(filePath):
                items[item] = getDirectoryStructure(filePath) # Recursively get the directory structure of the folder

    # Then, get all files and add them to the 'files' key of the folder dictionary
    items['files'] = []
    for item in os.listdir(rootDir):
        filePath = os.path.join(rootDir, item)

        if item[0] != '.':
            if os.path.isfile(filePath):
                items['files'].append(item)

    return items

def resourcePath(relativePath) -> str:
    ''' Replace the relative path with the absolute path to the resource. This allows files to be accessed when the program is compiled to an executable. '''
    try:
        basePath = sys._MEIPASS # PyInstaller creates a temp folder and stores path in _MEIPASS
    except Exception: # If not running as a PyInstaller executable, use the current directory
        basePath = os.path.abspath(".")

    return os.path.join(basePath, relativePath) # Return the absolute path to the resource

def showFileExplorer(file):
    ''' Function to open the file explorer to the file. '''
    filePath = resourcePath(file)
    if platform.system() == "Windows":
        #subprocess.Popen(f'explorer /select,"{filePath}"')
        folderPath = os.path.dirname(filePath)
        os.startfile(folderPath)
    elif platform.system() == "Darwin":
        subprocess.call(["open", "-R", filePath])
    else:
        subprocess.Popen(["xdg-open", filePath])

def copyFile(file):
    ''' Function to copy the file to the destination. In this case, the desktop. '''
    filePath = resourcePath(file)
    # destinationPath = resourcePath(destination)
    shutil.copy2(filePath, desktop)

def clearFolder(folder):
    ''' Function to clear the contents of a folder. '''
    folderPath = resourcePath(folder)
    for item in os.listdir(folderPath):
        itemPath = os.path.join(folderPath, item)
        if os.path.isfile(itemPath):
            os.unlink(itemPath)
        elif os.path.isdir(itemPath):
            shutil.rmtree(itemPath)

def sendEmail(receiver, message):
    ''' Function to send an email. '''
    port = 465
    smtpServer = "smtp.gmail.com"
    sender = 'agssoundandlighting@gmail.com'
    key = 'hrrs axmu thjx uzxd'

    with smtplib.SMTP_SSL(smtpServer, port) as server:
        server.login(sender, key)
        try:
            server.sendmail(sender, receiver, message)
            print("Successfully sent email")
        except Exception as e:
            print("Error: unable to send email")
            raise Exception()

def generateCode() -> int:
    ''' Function to generate a random 4 digit code. '''
    return random.randint(1000, 9999) # Return a random 4 digit code

def createRota(members: list):
    ''' Create a random 2 week rota in a Word document. '''
    def getRandomMembers(members) -> list:
        ''' Get 3 random members from the list of members. '''
        membersChosen = random.sample(members, 3)
        return '\n'.join(membersChosen)
    
    document = Document(resourcePath('Contents/Documents/Templates/AGS Sound and Lighting Assembly Rota Template.docx')) # open the template document

    styles = document.styles # get the styles from the document
    characterStyle = styles.add_style('commentsStyle', WD_STYLE_TYPE.PARAGRAPH) # add a character style to the document
    font = characterStyle.font # get the font from the character style
    font.name = 'EB Garamond' # set the font to Arial
    font.size = Pt(15) # set the font size to 0.2 inches

    stinsonHallA, stinsonHallB, sportsHallA, sportsHallB = [table for table in document.tables] # get the 4 tables from the document
    
    for day in range(5): # for each day in the week, add the members to that position in each of the tables
        stinsonHallApara = stinsonHallA.cell(0, 10+day)
        stinsonHallBpara = stinsonHallB.cell(0, 10+day)
        sportsHallApara = sportsHallA.cell(0, 10+day)
        sportsHallBpara = sportsHallB.cell(0, 10+day)
        
        for paragraph in [stinsonHallApara, stinsonHallBpara, sportsHallApara, sportsHallBpara]: # for each paragraph in the tables
            paragraph.text = f'{getRandomMembers(members)}'
            paragraph.paragraphs[0].style = characterStyle # set the font of the paragraph
            paragraph.paragraphs[0].alignment = 1 # set the alignment of the paragraph to center

    document.save(resourcePath('Contents/Documents/Current Working Documents/Rotas/AGS Sound and Lighting Assembly Rota.docx')) # save the document